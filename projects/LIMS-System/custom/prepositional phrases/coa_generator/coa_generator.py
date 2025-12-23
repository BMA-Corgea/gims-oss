# coa_generator.py — COA Generator (adapted to unified template)
# One file usable as a Prepositional Phrase (pphrase) under the new core.
# Edit SECTION A and SECTION B only. SECTION Z is glue; do not edit.

from __future__ import annotations
import os, json, csv, shutil, re
from pathlib import Path
from typing import Dict, List, Any, Optional
from collections import defaultdict
from datetime import datetime

# External dependency (python-docx) — used ONLY in host-side POST_DOC
try:
    from docx import Document
    from docx.shared import Pt, Mm
except Exception as _e:
    Document = None  # type: ignore
    Pt = Mm = None   # type: ignore

# =============================================================================
# SECTION A — DEV CONFIG (EDIT THIS)
# =============================================================================

TOOL_NAME: str = "COA_Generator"
TOOL_VERSION: str = "1.0.0"

# Choose mode: "parser" | "pphrase"
TOOL_KIND: str = "pphrase"

# RAW folder inputs (schema aliases). Backend ensures one file per folder and pre-digests if needed.
# For this pphrase, we scan the mounted run folders directly (discovered relative to the output folder),
# so we do not require explicit RAW mounts.
RAW_FOLDERS: List[str] = []

# Exact file inputs besides raw folders (schema names/aliases).
# Left empty; we locate Tests_log.jsonl and noun folders (COA Name Map, Primary Aromas, Submission)
# relative to the project structure at runtime.
FILE_INPUTS: List[str] = []

# Parser outputs (unused for pphrase)
PARSER_OUTPUT_FILES: List[str] = []

# Pphrase output: single folder name to write into (single segment, no slashes).
PPHRASE_OUTPUT_FOLDER: Optional[str] = None

# Optional DB inputs (pphrase)
DB_INPUTS: List[Dict[str, Any]] = [
    {"endpoint": "noun_items", "params": {"noun_type": "COA Name Map"},   "alias": "coa_name_map_items"},
    {"endpoint": "noun_items", "params": {"noun_type": "Primary Aromas"}, "alias": "primary_aromas_items"},
    {"endpoint": "noun_items", "params": {"noun_type": "Submission"},     "alias": "submission_items"},
    {"endpoint": "verb_group_log", "params": {"verb_group": "Tests"}, "alias": "tests_log"},
    {"endpoint": "verb_group_log_config", "params": {"verb_group": "Tests"}, "alias": "tests_log_config"},
    {"endpoint": "data_dump_dir", "params": {"verb_group": "Tests"}, "alias": "test_run_artifacts"},
]

# Host-side post-doc step (pphrase): run callable in THIS SAME MODULE
POST_DOC: Optional[Dict[str, Any]] = {
    "entry": "coa_generator:postdoc_render",
    "args": {
        "template_name": "Python COA Template.docx",
        "target_subfolder": None
    },
}

# ------- Pre-phrase UI schema (used when TOOL_KIND == "pphrase") -------
PREPHRASE_SETTINGS: List[Dict[str, Any]] = [
    {
        "id": "include_watermark",
        "label": "Include watermark",
        "kind": "bool",
        "default": True
    },
    {
        "id": "output_format",
        "label": "Output format",
        "kind": "single",
        "options": [
            {"label": "PDF",  "value": "pdf"},
            {"label": "DOCX", "value": "docx"}
        ],
        "default": "docx"
    },

    # SEPARATE FIELD FOR POTENCY SAMPLES
    {
        "id": "samples_potency",
        "label": "Potency Samples (complete)",
        "kind": "multi",
        "options": {
            "source": "noun: Potency Sample",
            "complete": True,
            "filters": [
                {"field": "Client", "op": "in", "ref": "clients"},
                {"field": "received_date", "op": "between", "ref": ["date_start", "date_end"]},
                {"field": "Sample ID", "op": "exists"},
                {"field": "Client",    "op": "exists"}
            ],
            "unique_by": ["Sample ID"],
            "map": {
                "label": "{Sample ID} — {Client} — {Sample Name} (Potency)",
                "value": "{Sample ID}"
            },
            "sort": [
                {"field": "Client",    "dir": "asc"},
                {"field": "Sample ID", "dir": "asc"}
            ],
            "limit": 1000
        },
        "default": []
    },

    # SEPARATE FIELD FOR TERPENE SAMPLES
    {
        "id": "samples_terpene",
        "label": "Terpene Samples (complete)",
        "kind": "multi",
        "options": {
            "source": "noun: Terpene Sample",
            "complete": True,
            "filters": [
                {"field": "Client", "op": "in", "ref": "clients"},
                {"field": "received_date", "op": "between", "ref": ["date_start", "date_end"]},
                {"field": "Sample ID", "op": "exists"},
                {"field": "Client",    "op": "exists"}
            ],
            "unique_by": ["Sample ID"],
            "map": {
                "label": "{Sample ID} — {Client} — {Sample Name} (Terpene)",
                "value": "{Sample ID}"
            },
            "sort": [
                {"field": "Client",    "dir": "asc"},
                {"field": "Sample ID", "dir": "asc"}
            ],
            "limit": 1000
        },
        "default": []
    }
]


# =============================================================================
# SECTION B — DEV WORK (EDIT THIS): define how inputs -> outputs
# =============================================================================

def work_parser(inputs: Dict[str, Any], outputs: Dict[str, str], params: Dict[str, Any]) -> None:
    # Not used for this tool; included to satisfy template. No-op.
    pass


def work_pphrase(inputs: Dict[str, Any], outputs: Dict[str, Any], params: Dict[str, Any]) -> None:
    """
    Container-side prep only:
      - use mounted inputs (noun items, run artifacts, tests_log)
      - resolve requested samples
      - consolidate potency/terpene runs per sample
      - compute minimal replacements payload (+ weights.json)
      - dump everything under OUTPUT_FOLDER/_prepared/** as JSON
    NO python-docx. NO template parsing here.
    """
    from pathlib import Path
    import json
    import shutil
    import csv

    def dbg(*a): print("[pphrase]", *a)

    # ------------------------
    # tiny helpers
    # ------------------------
    def _coerce_paths(v) -> list[Path]:
        if not v:
            return []
        if isinstance(v, (list, tuple, set)):
            return [Path(p) for p in v]
        return [Path(v)]

    def _get_any(d: dict, *keys):
        for k in keys:
            if k in d and d[k]:
                return d[k]
        return None

    def _gather_run_dirs(mounted: list[Path]) -> list[Path]:
        """
        Accepts a list of paths that may be:
          - direct per-run folders (…/data_dumps/{run_id}), or
          - a base folder containing many run folders (…/data_dumps)
        Returns the list of per-run folders that contain DataEntry.json.
        """
        out: list[Path] = []
        for p in mounted:
            if not p.exists():
                continue
            if (p / "DataEntry.json").exists():
                out.append(p)
            elif p.is_dir():
                # Treat as base; add children that look like run dirs
                for child in p.iterdir():
                    if child.is_dir() and (child / "DataEntry.json").exists():
                        out.append(child)
        # dedupe
        seen = set()
        uniq = []
        for rd in out:
            if rd.resolve() not in seen:
                uniq.append(rd)
                seen.add(rd.resolve())
        return uniq

    # ------------------------
    # OUTPUT root
    # ------------------------
    if "OUTPUT_FOLDER" not in outputs:
        dbg("No OUTPUT_FOLDER configured; skipping file output operations")
        return

    out_dir = Path(outputs["OUTPUT_FOLDER"]).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    prep_dir = out_dir / "_prepared"
    prep_dir.mkdir(parents=True, exist_ok=True)
    dbg("OUTPUT_FOLDER:", out_dir)

    # ------------------------
    # Mounted inputs
    # ------------------------
    coa_name_map_file      = Path(inputs["coa_name_map_items"])
    primary_aromas_file    = Path(inputs["primary_aromas_items"])
    submission_items_file  = Path(inputs["submission_items"])
    tests_log_file         = Path(inputs["tests_log"])
    tests_log_config_file  = Path(inputs["tests_log_config"])

    # Runs can be mounted either as a base dir or as a list of per-run dirs (GUI does the latter).
    # Try several likely aliases to be resilient to IoSpec alias naming.
    run_mount_value = _get_any(
        inputs,
        "test_run_artifacts",  # your original alias
        "data_dump_dir",       # alias defaults to endpoint name if not overridden
        "data_dumps",          # fallback guess
    )
    run_mounts = _coerce_paths(run_mount_value)
    run_dirs = _gather_run_dirs(run_mounts)
    dbg("mounted run path(s):", [str(p) for p in run_mounts] or "(none)")
    dbg("resolved run dir(s):", [p.name for p in run_dirs] or "(none)")

    # Bases (used for hints + fallback builder)
    nouns_base = submission_items_file.parent.parent  # .../nouns
    # If we have per-run dirs, use their parent as the runs base; else assume the single legacy base.
    runs_base = run_dirs[0].parent if run_dirs else (Path(run_mount_value) if run_mount_value else None)

    # Noun items (JSONL)
    def load_items_jsonl(path: Path):
        items = []
        if path.exists():
            with path.open(encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line:
                        items.append(json.loads(line))
        return items

    coa_name_map       = load_items_jsonl(coa_name_map_file)
    primary_aroma_itm  = load_items_jsonl(primary_aromas_file)
    submission_items   = load_items_jsonl(submission_items_file)
    dbg("COA Name Map items:", len(coa_name_map))
    dbg("Primary Aromas items:", len(primary_aroma_itm))
    dbg("Submission items:", len(submission_items))

    tests_config = {}
    if tests_log_config_file.exists():
        with tests_log_config_file.open(encoding="utf-8") as f:
            tests_config = json.load(f)
    dbg("Tests config loaded:", bool(tests_config))

    # ---- tolerant Tests_log reader
    def load_runid_meta(path: Path):
        """Return {run_id: {"date_tested": str|None, "test_type": str|None}}."""
        meta: Dict[str, dict] = {}
        if path.exists():
            with path.open(encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    rec = json.loads(line)
                    rid = (rec.get("run_ID") or rec.get("_runID") or
                           rec.get("runId") or rec.get("Run ID") or rec.get("run id"))
                    if not rid:
                        continue
                    date = (rec.get("date_tested") or rec.get("dateTested") or
                            rec.get("Date Tested") or rec.get("date"))
                    ttype = (rec.get("test_type") or rec.get("verb") or
                             rec.get("test") or rec.get("testType"))
                    meta[str(rid)] = {"date_tested": date, "test_type": ttype}
        return meta

    runid_meta  = load_runid_meta(tests_log_file)
    runid_dates = {rid: (m or {}).get("date_tested") for rid, m in runid_meta.items()}
    dbg("runid_meta loaded:", len(runid_meta))

    # Pre-scan for weights.csv per run (from the mounted run dirs).
    # Keys are run_id (assumed == run_dir.name), values are absolute file paths.
    weights_csv_map: Dict[str, str] = {}
    for rd in run_dirs:
        cand = rd / "weights.csv"
        if cand.exists():
            weights_csv_map[rd.name] = str(cand.resolve())
    dbg("weights.csv discovered for run(s):", list(weights_csv_map.keys()))

    # ---- Collect completed samples from the mounted run dirs
    def collect_completed_samples_from_run_dirs(_run_dirs: list[Path], _runid_meta: Dict[str, dict]):
        samples = []
        if not _run_dirs:
            dbg("No run directories were mounted/discovered")
            return samples

        for run_dir in _run_dirs:
            data_entry_file = run_dir / "DataEntry.json"
            if not data_entry_file.exists():
                continue

            try:
                raw = json.loads(data_entry_file.read_text(encoding="utf-8"))
            except Exception as e:
                dbg(f"Error reading {data_entry_file}: {e}")
                continue

            entries = raw if isinstance(raw, list) else [raw]
            for e in entries:
                sid    = (e.get("Sample ID") or e.get("sample_id") or e.get("sampleID"))
                sname  = (e.get("Sample Name") or e.get("sample_name"))
                stype  = (e.get("Sample Type") or e.get("sample_type"))
                client = (e.get("Client") or e.get("client"))
                rid    = (e.get("_runID") or e.get("run_ID") or e.get("runId") or e.get("Run ID") or e.get("run id"))
                if not sid or not rid:
                    continue

                meta         = _runid_meta.get(str(rid), {})
                ttype        = (e.get("test_type") or meta.get("test_type"))
                date_tested  = meta.get("date_tested")

                # heuristic if missing
                if not ttype:
                    name = run_dir.name.lower()
                    if "terp" in name:
                        ttype = "Terpene_Test"
                    elif "pote" in name or "canna" in name:
                        ttype = "Potency_Test"

                sample = {
                    "sample_id": sid,
                    "sample_name": sname,
                    "sample_type": stype,
                    "client": client,
                    "test_type": ttype,
                    "run_id": str(rid),
                    "date_tested": date_tested,
                    "data_entry_path": str(data_entry_file),
                    "run_dir": str(run_dir),
                }
                if e.get("potency"): sample["potency"] = e["potency"]
                if e.get("terpene"): sample["terpene"] = e["terpene"]

                samples.append(sample)
        return samples

    completed_samples = collect_completed_samples_from_run_dirs(run_dirs, runid_meta)
    dbg("completed_samples found:", len(completed_samples))

    # ------------------------
    # Filter requested samples
    # ------------------------
    requested_potency  = [str(s).strip() for s in (params.get("samples_potency") or []) if str(s).strip()]
    requested_terpene  = [str(s).strip() for s in (params.get("samples_terpene") or []) if str(s).strip()]

    # union of both lists
    requested_ids = sorted(set(requested_potency) | set(requested_terpene))
    dbg("requested sample IDs (potency ∪ terpene):", requested_ids)

    if not requested_ids:
        dbg("No sample IDs requested; aborting prep.")
        return

    # filter completed_samples down to those with a sample_id in the requested set
    filtered = [s for s in completed_samples if s.get("sample_id") in requested_ids]
    pre_count = len(filtered)

    # only keep Potency/Terpene runs
    filtered = [s for s in filtered if s.get("test_type") in ("Potency_Test", "Terpene_Test")]
    dbg(f"filtered for Potency/Terpene: {len(filtered)} (was {pre_count})")

    if not filtered:
        dbg("WARNING: no Potency/Terpene samples after filter; nothing to prepare")
        return

    # Consolidate per-sample (choose latest per test by date)
    def consolidate_sample_tests(samples):
        out: Dict[str, dict] = {}
        for s in samples:
            sid = s.get("sample_id")
            if not sid:
                continue
            rec = out.setdefault(sid, {
                "sample_id": sid,
                "sample_name": s.get("sample_name"),
                "sample_type": s.get("sample_type"),
                "client": s.get("client"),
                "potency": {},
                "terpene": {},
            })
            if s.get("test_type") == "Potency_Test":
                a = rec.get("potency") or {}
                if not a or (runid_dates.get(s["run_id"], "") or "") >= (runid_dates.get(a.get("run_id", ""), "") or ""):
                    rec["potency"] = s
            elif s.get("test_type") == "Terpene_Test":
                a = rec.get("terpene") or {}
                if not a or (runid_dates.get(s["run_id"], "") or "") >= (runid_dates.get(a.get("run_id", ""), "") or ""):
                    rec["terpene"] = s
        return list(out.values())

    samples_consolidated = consolidate_sample_tests(filtered)
    dbg("consolidated samples:", len(samples_consolidated))

    # ------------------------
    # Build + write payloads per sample
    # ------------------------
    # For the post-doc step, also publish a run_id -> weights.csv map at the top level (once).
    if weights_csv_map:
        (prep_dir / "weights_csv_map.json").write_text(json.dumps(weights_csv_map, indent=2), encoding="utf-8")

    for sample in samples_consolidated:
        client    = sample.get("client") or "UnknownClient"
        sample_id = sample.get("sample_id") or "UnknownSample"

        pot_info  = sample.get("potency", {}) or {}
        ter_info  = sample.get("terpene", {}) or {}

        replacements = {
            "Client": str(client),
            "sampleID": str(sample_id),  # Force to string to prevent formatting
            "sampleName": str(sample.get("sample_name", "n/a")),
            "sampleType": str(sample.get("sample_type", "n/a")),
            "terpeneDateTested": str(ter_info.get("date_tested", "n/a")),  # Keep dates as strings
            "cannabinoidDateTested": str(pot_info.get("date_tested", "n/a")),  # Keep dates as strings
        }

        # Prefer mounted weights.csv paths if available; otherwise let the existing helper build weights.json.
        pot_run_id  = pot_info.get("run_id")
        terp_run_id = ter_info.get("run_id")
        pot_csv     = weights_csv_map.get(str(pot_run_id)) if pot_run_id else None
        terp_csv    = weights_csv_map.get(str(terp_run_id)) if terp_run_id else None

        if runs_base is None and run_dirs:
            runs_base = run_dirs[0].parent  # final fallback

        # Keep your original JSON structure (build_weights_for_sample), but also expose CSVs when present.
        weights_data = None
        try:
            weights_data = build_weights_for_sample(
                {
                    "sample_id": sample_id,
                    "potency": {"run_id": pot_run_id} if pot_run_id else {},
                    "terpene": {"run_id": terp_run_id} if terp_run_id else {},
                },
                runs_base
            )
        except Exception as e:
            dbg("build_weights_for_sample failed; continuing without it:", e)

        sdir = prep_dir / client / sample_id
        sdir.mkdir(parents=True, exist_ok=True)

        # Write core files
        (sdir / "replacements.json").write_text(json.dumps(replacements, indent=2), encoding="utf-8")
        if weights_data is not None:
            (sdir / "weights.json").write_text(json.dumps(weights_data, indent=2), encoding="utf-8")

        # If we have CSVs, copy them next to the JSON for convenience and advertise in context.
        weights_csv_paths = {}
        if pot_csv and Path(pot_csv).exists():
            dst = sdir / "weights_potency.csv"
            try:
                shutil.copy2(pot_csv, dst)
                weights_csv_paths["potency"] = str(dst)
            except Exception as e:
                dbg("copy weights_potency.csv failed:", e)
                weights_csv_paths["potency"] = str(pot_csv)  # fall back to original absolute path
        if terp_csv and Path(terp_csv).exists():
            dst = sdir / "weights_terpene.csv"
            try:
                shutil.copy2(terp_csv, dst)
                weights_csv_paths["terpene"] = str(dst)
            except Exception as e:
                dbg("copy weights_terpene.csv failed:", e)
                weights_csv_paths["terpene"] = str(terp_csv)

        # Per-sample context for host-side postdoc
        (sdir / "context.json").write_text(json.dumps({
            # directory hints for the host-side postdoc
            "nouns_base": str(nouns_base),
            "runs_base": str(runs_base) if runs_base else None,
            "primary_aromas_dir": str(primary_aromas_file.parent),
            "submission_dir": str(submission_items_file.parent),
            "tests_log_path": str(tests_log_file),

            # NEW: expose mounted run dir + weights CSV locations the runner provided
            "run_dir": sample.get("run_dir"),
            "weights_csv_paths": weights_csv_paths,                 # copied into sdir if possible; else absolute mounts
            "weights_csv_map": weights_csv_map,                    # all runs in this execution (top-level map also written)
        }, indent=2), encoding="utf-8")

        dbg("WROTE:", sdir / "replacements.json")

    dbg(f"prep complete → {prep_dir}")
    dbg(f"processed {len(samples_consolidated)} consolidated samples")

# ---------------- Host-side POST_DOC callable (same module) -----------------

def postdoc_render(
    env: Dict[str, Any],
    *,
    template_name: str = "Python COA Template.docx",
    target_subfolder: Optional[str] = None,
    output_root: str,                 # required (runner injects the allowed root)
    phrase_root: Optional[str] = None,
    project_path: Optional[str] = None,
    context: Optional[dict] = None,
    **kwargs,
) -> str:
    """
    Host-side: read OUTPUT_FOLDER/_prepared/** and generate DOCX files.
    Template is read from the phrase folder; outputs are saved under:
    {output_root}/{client}/{YYYY-MM-DD}/COA for {Sample ID}.docx
    Also emits a JSON manifest listing all created files (absolute paths).
    """
    from datetime import datetime as _dt
    require_docx()

    # prefer injected output_root; runner enforces writes under this anchor
    if not output_root or "tmp" in str(output_root):
        output_root = (
            kwargs.get("output_root")
            or env.get("prepositional_phrase_output_dir")
            or (env.get("outputs") or {}).get("OUTPUT_FOLDER")
            or env.get("work_dir", "")
        )

    ctx_env = env.get("context", {}) or {}
    # Default: template lives alongside this module
    phrase_root = Path(kwargs.get("phrase_root") or Path(__file__).parent)

    out_root = Path(output_root)
    prep_dir = out_root / "_prepared"

    include_watermark = bool(((ctx_env or {}).get("params") or {}).get("include_watermark"))

    # Resolve template
    template = phrase_root / template_name
    if not template.exists():
        raise FileNotFoundError(f"Template not found at {template}")

    today_date = _dt.now().strftime("%Y-%m-%d")

    # Discover prepared samples
    prepared_entries = []
    for repl_path in list(prep_dir.glob("*/*/replacements.json")) + list(prep_dir.glob("*/replacements.json")):
        sdir = repl_path.parent
        client_name = sdir.parent.name if sdir.parent.name != "_prepared" else sdir.name
        replacements = _load_json(sdir / "replacements.json") or {}
        weights_data = _load_json(sdir / "weights.json") or {}
        ctx = _load_json(sdir / "context.json") or {}
        sample_id = replacements.get("sampleID") or sdir.name
        prepared_entries.append({
            "sample_id": sample_id,
            "client_name": client_name,
            "replacements": replacements,
            "weights_data": weights_data,
            "context": ctx,
        })

    # Track what we create for the manifest
    created_docs: List[Dict[str, Any]] = []

    # Generate a DOCX per prepared sample
    for entry in prepared_entries:
        sample_id     = entry["sample_id"]
        client_name   = entry["client_name"]
        replacements  = entry["replacements"]
        weights_data  = entry["weights_data"]
        ctx           = entry["context"] or {}

        nouns_base = Path(ctx.get("nouns_base")) if ctx.get("nouns_base") else None
        runs_base  = Path(ctx.get("runs_base"))  if ctx.get("runs_base")  else None
        primary_aromas_dir = Path(ctx.get("primary_aromas_dir")) if ctx.get("primary_aromas_dir") else None
        submission_dir     = Path(ctx.get("submission_dir"))     if ctx.get("submission_dir") else None
        tests_log_path     = Path(ctx.get("tests_log_path"))     if ctx.get("tests_log_path") else None

        # Create directory structure: {output_root}/{client}/{YYYY-MM-DD}/
        client_date_dir = out_root / client_name / today_date
        client_date_dir.mkdir(parents=True, exist_ok=True)

        # Create docx from template
        dst = client_date_dir / f"COA for {sample_id}.docx"
        doc = Document(str(template))

        # -----------------------------
        # 1) Simple placeholder replace
        #    (avoid formatting IDs/dates)
        # -----------------------------
        no_format_keys = {"sampleID", "sampleName", "sampleType",
                          "terpeneDateTested", "cannabinoidDateTested", "sampleReceived", "Client"}
        repl_pairs = [
            (f"{{{{{k}}}}}", v if k in no_format_keys else _safe_format(v))
            for k, v in (replacements or {}).items()
        ]
        replace_placeholders_in_docx(doc, repl_pairs)

        # ------------------------------------------------
        # 2) Two-step placeholders (amountTHCA/D9/CBD/etc)
        # ------------------------------------------------
        if nouns_base is not None:
            coa_name_map_path = nouns_base / "COA Name Map" / "items.jsonl"
        else:
            coa_name_map_path = None  # type: ignore

        two_step_pairs = compute_two_step_placeholders(weights_data, coa_name_map_path) if coa_name_map_path else []
        if two_step_pairs:
            replace_placeholders_in_docx(doc, two_step_pairs)

        # ------------------------------------------------
        # 3) Tables & visuals
        # ------------------------------------------------
        # Build analyte table (requires nouns_base for name map)
        if nouns_base is not None:
            coa_name_map_path = nouns_base / "COA Name Map" / "items.jsonl"
            if coa_name_map_path.exists() and weights_data:
                analyte_table = build_analyte_table_data(weights_data, coa_name_map_path, template)
                insert_analyte_table_data(doc, analyte_table)

        # Primary aromas (top 3) and images — uses nouns_base
        if nouns_base is not None and primary_aromas_dir and (primary_aromas_dir / "items.jsonl").exists():
            coa_name_map_items = load_items_jsonl(nouns_base / "COA Name Map" / "items.jsonl")
            primary_aroma_items = load_items_jsonl(primary_aromas_dir / "items.jsonl")
            insert_primary_aromas(
                doc=doc,
                weights_data=weights_data,
                coa_name_map=coa_name_map_items,
                primary_aroma_pictures=primary_aroma_items,
                nouns_base=nouns_base,
            )

        # Submission date/image — needs both runs_base (to find DataEntry) and nouns_base (for images)
        if nouns_base is not None and runs_base is not None and submission_dir and tests_log_path and tests_log_path.exists():
            insert_sample_submission_data(
                doc=doc,
                sample_id=sample_id,
                submissions_path=submission_dir,
                runs_base=runs_base,
                nouns_base=nouns_base,
                tests_log_path=tests_log_path,
            )

        # Optional watermark
        if include_watermark:
            for section in doc.sections:
                footer = section.footer
                p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                if "DRAFT" not in p.text:
                    p.text = (p.text + "  ·  DRAFT WATERMARK").strip()

        doc.save(str(dst))
        dbg(f"is this thing on? can anyone hear me?????????????????????????????????????")
        dbg(f"[postdoc] saved COA → {dst}")

        created_docs.append({
            "sample_id": sample_id,
            "client": client_name,
            "path": str(dst.resolve()),
            "created_at": _dt.now().isoformat(timespec="seconds"),
            "format": "docx",
        })

    # ─────────────────────────────────────────────────────────────────────
    # DO NOT DELETE — EMIT MANIFEST FOR DOWNSTREAM PICKUP (GUI / SYNC JOBS)
    # This block writes a small JSON file that lists every file this postdoc
    # just created (with absolute paths). Several parts of the system (e.g.,
    # GUI download buttons, S3 syncers, notifications) look for this manifest.
    # If you remove or rename it, those features will silently break.
    # ─────────────────────────────────────────────────────────────────────
    try:
        manifest = {
            "note": (
                "This manifest is generated by postdoc_render to declare the files "
                "it created on the host filesystem. Downstream tooling consumes this. "
                "Do not delete casually."
            ),
            "generator": {
                "module": __name__,
                "function": "postdoc_render",
                "template": str(template.resolve()),
                "version": TOOL_VERSION,
            },
            "context": {
                "pphrase_name": ctx_env.get("pphrase_name"),
                "project_path": project_path or env.get("project_path"),
                "output_root": str(out_root.resolve()),
                "executed_at": _dt.now().isoformat(timespec="seconds"),
            },
            "created": created_docs,
        }
        # Write a stable filename (always overwritten) and a timestamped snapshot.
        stable_path = out_root / "_postdoc_outputs.json"
        snapshot_path = out_root / f"_postdoc_outputs_{_dt.now().strftime('%Y%m%d_%H%M%S')}.json"

        stable_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
        snapshot_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

        dbg(f"[postdoc] wrote manifest → {stable_path}")
        dbg(f"[postdoc] wrote manifest snapshot → {snapshot_path}")
    except Exception as _e:
        dbg(f"[postdoc][warn] failed to write manifest: {_e}")
    # ─────────────────────────────────────────────────────────────────────

    return str(out_root)


# =============================================================================
# SECTION Z — BACKEND GLUE (DO NOT EDIT BELOW)
# =============================================================================

TOOL = {
    "name": TOOL_NAME,
    "version": TOOL_VERSION,
    "kind": TOOL_KIND,
    "about": (
        "Parses raw outputs to produce interpretation files."
        if TOOL_KIND == "parser"
        else "Generates a document/package from run artifacts."
    ),
}

def get_metadata() -> dict:
    return dict(TOOL)

def get_PREPHRASE_SETTINGS() -> List[Dict[str, Any]]:
    return list(PREPHRASE_SETTINGS)

def get_io_spec() -> dict:
    if TOOL["kind"] not in ("parser", "pphrase"):
        raise ValueError("TOOL.kind must be 'parser' or 'pphrase'")

    spec: Dict[str, Any] = {
        "kind": TOOL["kind"],
        "raw_folders": list(RAW_FOLDERS),
        "file_inputs": list(FILE_INPUTS),
    }
    if TOOL["kind"] == "parser":
        spec["outputs"] = {"files": list(PARSER_OUTPUT_FILES)}
    else:
        spec["outputs"] = {"folder": PPHRASE_OUTPUT_FOLDER or ""}
        extra: Dict[str, Any] = {}
        if DB_INPUTS:
            extra["db_inputs"] = list(DB_INPUTS)
        if POST_DOC:
            extra["post_doc"] = dict(POST_DOC)
        if extra:
            spec["extra"] = extra
    return spec

# ---- tiny IO helpers for dev section ----
def _read_csv(path: str) -> List[Dict[str, Any]]:
    with open(path, "r", newline="") as f:
        return [row for row in csv.DictReader(f)]

def _write_csv(path: str | Path, rows: List[Dict[str, Any]]) -> None:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        path.write_text("")
        return
    with open(path, "w", newline="") as f:
        w = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        w.writeheader()
        w.writerows(rows)

# ---- runtime entrypoint (sandbox calls this) ----
def run() -> None:
    mapping = os.environ.get("GIMS_IO_JSON")
    if not mapping:
        raise RuntimeError("GIMS_IO_JSON not provided")
    io_map = json.loads(mapping)

    kind = io_map.get("kind")
    inputs = io_map.get("inputs", {})
    outputs = io_map.get("outputs", {})
    params = io_map.get("params", {}) or {}

    # Normalize: ensure raw folder values are lists for dev ergonomics
    norm_inputs: Dict[str, Any] = {}
    for k, v in inputs.items():
        if isinstance(v, list):
            norm_inputs[k] = v
        elif isinstance(v, str):
            norm_inputs[k] = v
        else:
            norm_inputs[k] = v

    if kind == "parser":
        work_parser(norm_inputs, outputs, params)
    elif kind == "pphrase":
        work_pphrase(norm_inputs, outputs, params)
    else:
        raise RuntimeError(f"Unknown tool kind: {kind}")

if __name__ == "__main__":
    run()

# =============================================================================
# ------------------------- Helper functions (internal) -----------------------
# =============================================================================

def dbg(*args):
    """Lightweight debug logger (always on for this tool)."""
    print(*args)

def require_docx():
    if Document is None:
        raise RuntimeError("python-docx is required for COA generation on the host, but is not available.")

def find_nearest(start: Path, target_name: str) -> Optional[Path]:
    start = Path(start).resolve()
    for ancestor in [start] + list(start.parents):
        candidate = ancestor / target_name
        if candidate.exists():
            return candidate
    return None

def _load_json(p: Path) -> Any:
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else None

def _safe_format(val):
    """
    Format numeric *values* only. Leave plain strings like '957' or '2025-08-21' alone.
    Rules:
      - If val is int/float -> format 0 => 'ND', else 2 decimals
      - If val is a string that contains '.' or scientific notation -> try numeric format
      - Otherwise return as-is (prevents '957' → '957.00', dates → '2025-08-21.00')
    """
    # True numerics: format
    if isinstance(val, (int, float)):
        num = float(val)
        return "ND" if num == 0 else f"{num:.2f}"

    # Strings: only treat as numeric if they *look* like decimals/scientific
    s = str(val)
    if any(c in s for c in ".eE"):
        try:
            num = float(s)
            return "ND" if num == 0 else f"{num:.2f}"
        except (ValueError, TypeError):
            pass

    # Otherwise, preserve exactly
    return s

def load_items_jsonl(path: Path) -> List[dict]:
    items: List[dict] = []
    try:
        if path and path.exists():
            with path.open("r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line:
                        try:
                            items.append(json.loads(line))
                        except Exception as e:
                            dbg("[items.jsonl] bad line:", e, "|", line[:120], "...")
        else:
            dbg("[items.jsonl] not found:", path)
    except Exception as e:
        dbg("[items.jsonl] error:", e)
    return items

def load_runid_dates(tests_log_path: Path) -> Dict[str, str]:
    # legacy helper (kept for compatibility; unused by new prep)
    runid_dates: Dict[str, str] = {}
    try:
        with open(tests_log_path, "r", encoding="utf-8") as f:
            for line in f:
                entry = json.loads(line)
                rid = entry.get("run_ID")
                dt  = entry.get("date_tested")
                if rid:
                    runid_dates[str(rid)] = dt
    except FileNotFoundError:
        dbg("[runid_dates] Tests_log.jsonl not found at", tests_log_path)
    except Exception as e:
        dbg("[runid_dates] error:", e)
    return runid_dates

def build_weights_for_sample(sample: dict, inputs_base: Path) -> dict:
    """
    For the given consolidated sample record, load Weights.csv for
    Terpene_Test and Potency_Test and compute totals.
    inputs_base -> runs_base (data_dumps)
    """
    combined = {"sample": sample.get("sample_id")}

    def load_weights(test_key: str, verb_name: str):
        info = sample.get(test_key)
        if not info:
            return
        rid = info.get("run_id")
        if not rid:
            return
        folder = inputs_base / f"{rid}"
        weights_csv = folder / "Weights.csv"
        if not weights_csv.exists():
            dbg(f"[weights] missing: {weights_csv}")
            return

        rows = []
        with weights_csv.open("r", newline="", encoding="utf-8") as f:
            reader = csv.reader(f)
            rows = list(reader)

        if not rows:
            dbg(f"[weights] empty: {weights_csv}")
            return

        headers = rows[0]
        # find the row for our sample_id
        sample_row = None
        for row in rows[1:]:
            if row and row[0].strip() == sample.get("sample_id"):
                sample_row = row
                break
        if not sample_row:
            dbg(f"[weights] Sample ID {sample.get('sample_id')} not found in {weights_csv}")
            return

        data_dict = dict(zip(headers[1:], sample_row[1:]))
        total = 0.0
        for k, v in list(data_dict.items()):
            try:
                num = float(v)
                data_dict[k] = num
                total += num
            except (ValueError, TypeError):
                pass
        data_dict[f"{test_key}_total"] = round(total, 3)
        combined[test_key] = data_dict

    load_weights("terpene", "Terpene_Test")
    load_weights("potency", "Potency_Test")

    dbg("[weights] built:", combined)
    return combined

def replace_placeholders_in_docx(doc: Document, replacements: List[tuple[str, str]]):
    def replace_in_paragraph(paragraph, replacements):
        for placeholder, replacement in replacements:
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)

    def replace_in_table(table, replacements):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
                for nested_table in cell.tables:
                    replace_in_table(nested_table, replacements)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        replace_in_table(table, replacements)

def extract_analytes_by_table(doc_path: Path, display_to_entry: dict) -> List[dict]:
    doc = Document(str(doc_path))
    rows: List[dict] = []

    def parse_table(table, index):
        for tr in table.rows:
            name = tr.cells[0].text.strip()
            if name and (name in display_to_entry or name.lower() == "total"):
                entry = display_to_entry.get(name, {})
                rows.append({
                    "table_index": index,
                    "analyte": name,
                    "instrument_code": entry.get("Instrument Code"),
                    "analyte_type": entry.get("Analyte Type"),
                })
            for cell in tr.cells:
                for nested in cell.tables:
                    parse_table(nested, f"{index}-nested")

    for i, tbl in enumerate(doc.tables):
        parse_table(tbl, str(i))

    return rows

def build_analyte_table_data(weights_data: dict, coa_name_map_path: Path, coa_template_path: Path) -> List[dict]:
    with open(coa_name_map_path, "r", encoding="utf-8") as f:
        name_map = [json.loads(line) for line in f if line.strip()]
    display_to_entry = {e["COA Display Name"]: e for e in name_map}

    extracted = extract_analytes_by_table(coa_template_path, display_to_entry)

    type_to_key = {"Terpene": "terpene", "Cannabinoid": "potency"}
    last_type = None

    for row in extracted:
        analyte = row["analyte"]

        if analyte.lower() == "total":
            if last_type:
                key = type_to_key.get(last_type)
                total_key = f"{key}_total"
                raw_tot = weights_data.get(key, {}).get(total_key)
                if raw_tot is None:
                    pct, mg = "NT", "NT"
                else:
                    val = float(raw_tot)
                    pct = val
                    mg  = round(val * 10, 3)
            else:
                pct, mg = "NT", "NT"
        else:
            atype = row["analyte_type"]
            code  = row["instrument_code"]
            last_type = atype
            key = type_to_key.get(atype)
            raw = weights_data.get(key, {}).get(code)

            if raw is None or raw == "NaN":
                pct, mg = "NT", "NT"
            else:
                try:
                    val = float(raw)
                    if val == 0:
                        pct, mg = "ND", "ND"
                    else:
                        pct = val
                        mg  = round(val * 10, 3)
                except Exception:
                    pct, mg = "NT", "NT"

        row["%"]    = pct
        row["mg/g"] = mg

    dbg("[analytes] built rows:", len(extracted))
    return extracted

def set_cell_font_size(cell, size_pt):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)

def insert_analyte_table_data(doc: Document, analyte_table_data: List[dict]):
    for row_data in analyte_table_data:
        table_index = row_data["table_index"]
        analyte = row_data["analyte"]
        percent = row_data["%"]
        mg_per_g = row_data["mg/g"]

        if isinstance(table_index, str) and "-nested" in table_index:
            main_index, _ = table_index.split("-nested")
            table = doc.tables[int(main_index)]

            for row in table.rows:
                for cell in row.cells:
                    for nested_table in cell.tables:
                        for t_row in nested_table.rows:
                            if t_row.cells and t_row.cells[0].text.strip() == analyte:
                                if len(t_row.cells) >= 3:
                                    t_row.cells[1].text = str(percent)
                                    t_row.cells[2].text = str(mg_per_g)
                                    set_cell_font_size(t_row.cells[1], 8)
                                    set_cell_font_size(t_row.cells[2], 8)
                                elif len(t_row.cells) == 2:
                                    t_row.cells[1].text = str(percent)
                                    set_cell_font_size(t_row.cells[1], 8)
                                dbg(f"[analytes] filled nested {analyte}: %={percent}, mg/g={mg_per_g}")
        else:
            table = doc.tables[int(table_index)]
            for t_row in table.rows:
                if t_row.cells and t_row.cells[0].text.strip() == analyte:
                    if len(t_row.cells) >= 3:
                        t_row.cells[1].text = str(percent)
                        t_row.cells[2].text = str(mg_per_g)
                        set_cell_font_size(t_row.cells[1], 8)
                        set_cell_font_size(t_row.cells[2], 8)
                    elif len(t_row.cells) == 2:
                        t_row.cells[1].text = str(percent)
                        set_cell_font_size(t_row.cells[1], 8)
                    dbg(f"[analytes] filled {analyte}: %={percent}, mg/g={mg_per_g}")

def insert_primary_aromas(
    doc: Document,
    weights_data: dict,
    coa_name_map: List[dict],
    primary_aroma_pictures: List[dict],
    nouns_base: Path,
):
    """
    Inserts top unique primary aromas based on highest terpene concentrations,
    replacing both aroma and picture placeholders in-place (top 3).
    """
    instrument_to_aroma = {
        entry.get("Instrument Code"): entry.get("Primary Aroma")
        for entry in coa_name_map
        if entry.get("Analyte Type") == "Terpene" and entry.get("Primary Aroma")
    }

    aroma_to_path: Dict[str, Path] = {}
    for entry in primary_aroma_pictures:
        aroma = entry.get("Aroma")
        raw_id = entry.get("Picture ID")
        if aroma and raw_id:
            rel_path = raw_id.split("/", 1)[1] if raw_id.startswith("nouns/") else raw_id
            aroma_to_path[aroma] = nouns_base / rel_path

    terpene_data = weights_data.get("terpene", {})
    terpene_concs = []
    for code, val in terpene_data.items():
        if code in instrument_to_aroma and val not in ("NaN", "NT"):
            try:
                terpene_concs.append((code, float(val)))
            except Exception:
                pass
    top3 = sorted(terpene_concs, key=lambda x: x[1], reverse=True)[:3]

    unique_aromas: List[str] = []
    for code, _v in top3:
        aroma = instrument_to_aroma.get(code)
        if aroma and aroma not in unique_aromas:
            unique_aromas.append(aroma)
        if len(unique_aromas) >= 3:
            break

    aroma_repl = []
    for i in range(1, 4):
        aroma_ph = f"{{{{primaryAroma{i}}}}}"
        aroma_text = unique_aromas[i - 1] if i <= len(unique_aromas) else ""
        aroma_repl.append((aroma_ph, aroma_text))
    replace_placeholders_in_docx(doc, aroma_repl)

    for i in range(1, 4):
        picture_ph = f"{{{{primaryPicture{i}}}}}"
        aroma_text = unique_aromas[i - 1] if i <= len(unique_aromas) else ""
        pic_path = aroma_to_path.get(aroma_text)

        for p in doc.paragraphs:
            replace_placeholder_across_runs(p.runs, picture_ph, pic_path, 20, 20)
        for tbl in doc.tables:
            process_table(tbl, picture_ph, pic_path, 20, 20)

def replace_placeholder_across_runs(runs, placeholder, pic_path: Optional[Path], width_mm=None, height_mm=None):
    full_text = "".join(run.text for run in runs)
    if placeholder not in full_text:
        return False

    remaining = placeholder
    for run in runs:
        if not remaining:
            break
        if remaining in run.text:
            run.text = run.text.replace(remaining, "")
            remaining = ""
        elif remaining.startswith(run.text):
            remaining = remaining[len(run.text):]
            run.text = ""
        else:
            pass

    if pic_path and pic_path.exists():
        w = Mm(width_mm if width_mm is not None else 50)
        h = Mm(height_mm if height_mm is not None else 50)
        runs[0].add_picture(str(pic_path), width=w, height=h)
    return True

def process_cell(cell, picture_ph, pic_path, width_mm, height_mm):
    for p in cell.paragraphs:
        replace_placeholder_across_runs(p.runs, picture_ph, pic_path, width_mm, height_mm)
    for tbl in cell.tables:
        process_table(tbl, picture_ph, pic_path, width_mm, height_mm)

def process_table(tbl, picture_ph, pic_path, width_mm, height_mm):
    for row in tbl.rows:
        for cell in row.cells:
            process_cell(cell, picture_ph, pic_path, width_mm, height_mm)

# ---- DB helpers for noun instances / Picture adjectives --------------------
_SANITIZE_RE = re.compile(r"[^0-9a-zA-Z_]")

def _sanitize_table_name(noun_type: str) -> str:
    """
    Mirror noun_workbench: noun name → SQL table name (SQLite/RDS).
    Example: 'Submission' -> 'noun_Submission'
    """
    base = _SANITIZE_RE.sub("_", noun_type).strip("_")
    if not base or not base[0].isalpha():
        base = f"T_{base}"
    return f"noun_{base}"

def _prefixed(project: str, noun_table: str) -> str:
    """
    Mirror noun_workbench Postgres naming: <project>_noun_<Sanitized>.
    """
    base = noun_table
    if base.startswith("noun_"):
        base = base[len("noun_"):]
    return f"{project}_noun_{base}"

def _open_noun_db(project_path: Path):
    """
    Resolve object_sql_db via resolver and open a connection.

    Returns (kind, conn) where kind ∈ {'pg', 'sqlite', 'none'}.
    This function is only called host-side (postdoc), so it's safe
    to import api.* lazily here.
    """
    try:
        from api.manifest.resolver import resolve_path, get_db_uri  # type: ignore
    except Exception as e:
        dbg("[noun-db] resolver import failed; cannot open DB:", e)
        return "none", None

    uri = None
    try:
        uri = get_db_uri("object_sql_db")
        dbg("[noun-db] get_db_uri('object_sql_db') →", uri)
    except Exception as e:
        dbg("[noun-db] get_db_uri('object_sql_db') failed; will try SQLite fallback:", e)

    # Postgres first, if configured
    if uri and str(uri).startswith("postgresql"):
        try:
            import psycopg  # type: ignore
            dsn = str(uri)
            if dsn.startswith("postgresql+"):
                dsn = "postgresql://" + dsn.split("postgresql+", 1)[1]
            if "?ssl=require" in dsn:
                dsn = dsn.replace("?ssl=require", "?sslmode=require")
            dsn = dsn.replace("postgresql://asyncpg://", "postgresql://")
            dbg("[noun-db] connecting PG:", dsn)
            conn = psycopg.connect(dsn, autocommit=False)
            return "pg", conn
        except Exception as e:
            dbg("[noun-db] psycopg connect failed; falling back to SQLite:", e)

    # SQLite fallback via resolver
    try:
        from api.manifest.resolver import resolve_path  # type: ignore
        db_path = resolve_path(project_path, "object_sql_db")
    except Exception as e:
        dbg("[noun-db] resolve_path('object_sql_db') failed:", e)
        return "none", None

    import sqlite3
    dbg("[noun-db] connecting SQLite:", db_path)
    conn = sqlite3.connect(db_path.as_posix())
    conn.row_factory = sqlite3.Row
    return "sqlite", conn

def _load_noun_instance(project_path: Path, noun_type: str, primary_val: str) -> Optional[dict]:
    """
    Load a single noun instance from object_sql_db by primary ID.
    Uses the same table naming conventions as noun_workbench, including
    meta_tables overrides on Postgres when available.
    """
    if not primary_val:
        return None

    try:
        from api.i_o import get_noun_schema  # type: ignore
    except Exception as e:
        dbg("[noun-db] cannot import get_noun_schema; skipping DB lookup:", e)
        return None

    schema = get_noun_schema(project_path, noun_type) or {}
    primary_field = (schema or {}).get("primary_id_field") or "id"

    kind, conn = _open_noun_db(project_path)
    if not conn or kind == "none":
        return None

    project_name = project_path.name

    try:
        if kind == "pg":
            cur = conn.cursor()
            try:
                table_name = _prefixed(project_name, _sanitize_table_name(noun_type))

                # Try meta_tables override if present
                try:
                    cur.execute("""
                        SELECT table_name, primary_id
                        FROM meta_tables
                        WHERE project = %s AND noun_name = %s
                        LIMIT 1
                    """, (project_name, noun_type))
                    row = cur.fetchone()
                    if row:
                        tname, pid = row
                        if tname:
                            table_name = tname
                        if pid:
                            primary_field = pid
                except Exception as e:
                    dbg("[noun-db] meta_tables lookup failed; using fallback table name:", e)

                try:
                    cur.execute(f'SELECT * FROM "{table_name}" WHERE "{primary_field}" = %s LIMIT 1', (primary_val,))
                    row = cur.fetchone()
                except Exception as e:
                    dbg(f"[noun-db] PG query failed for table {table_name}: {e}")
                    return None

                if row is None:
                    return None

                cols = [c.name for c in cur.description]
                return dict(zip(cols, row))
            finally:
                cur.close()

        elif kind == "sqlite":
            import sqlite3
            table_name = _sanitize_table_name(noun_type)
            try:
                cur = conn.execute(f'SELECT * FROM "{table_name}" WHERE "{primary_field}" = ?', (primary_val,))
            except Exception as e:
                dbg(f"[noun-db] SQLite query failed for table {table_name}: {e}")
                return None
            row = cur.fetchone()
            if row is None:
                return None
            return dict(row)

    finally:
        try:
            conn.close()
        except Exception:
            pass

    return None

def _get_picture_field_name(project_path: Path, noun_type: str) -> Optional[str]:
    """
    Look up which field on this noun is a Picture adjective, based on noun_types.json.
    """
    try:
        from api.i_o import get_noun_schema  # type: ignore
    except Exception as e:
        dbg("[noun-db] cannot import get_noun_schema (Picture field):", e)
        return None

    schema = get_noun_schema(project_path, noun_type) or {}
    fields = (schema or {}).get("fields", {}) or {}
    for fname, fmeta in fields.items():
        if not isinstance(fmeta, dict):
            continue
        if fmeta.get("type") != "adjective":
            continue
        adj_class = (fmeta.get("adjective_class") or fmeta.get("class") or "").strip()
        if adj_class == "Picture":
            return fname
    return None

def _load_noun_picture_value(project_path: Path, noun_type: str, primary_val: str) -> Optional[str]:
    """
    Resolve the Picture adjective value for a given noun instance, if any.
    """
    field_name = _get_picture_field_name(project_path, noun_type)
    if not field_name:
        return None
    row = _load_noun_instance(project_path, noun_type, primary_val)
    if not row:
        return None
    raw = row.get(field_name)
    if raw is None:
        return None
    s = str(raw).strip()
    return s or None

def _load_submission_received_date(project_path: Path, submission_id: str) -> Optional[str]:
    """
    Convenience helper: pull received_date from Submission noun instance.
    """
    row = _load_noun_instance(project_path, "Submission", submission_id)
    if not row:
        return None

    # Prefer canonical key; fall back to a few common variants
    for key in ("received_date", "Received_Date", "Received Date", "receivedDate"):
        val = row.get(key)
        if val:
            return str(val).strip()

    val = row.get("received_date")
    if not val:
        return None
    return str(val).strip()

def insert_sample_submission_data(
    doc: Document,
    sample_id: str,
    submissions_path: Path,
    runs_base: Path,
    nouns_base: Path,
    tests_log_path: Path,
):
    """
    Inserts sample received date and submission/sample image into the COA document.

    New behavior (DB-aware, RDS/SQLite):
      - Uses object_sql_db (via resolver) instead of JSONL to fetch:
          * Submission.received_date
          * Picture adjectives on:
              - Potency Sample / Terpene Sample (preferred when available)
              - Sample (generic fallback)
              - Submission (last fallback)
      - Still uses runs_base to locate the correct DataEntry.json for the sample.
      - Picture values are resolved relative to nouns_base, respecting 'nouns/...' prefixes.
    """
    dbg("[submission] start for sample:", sample_id)

    # ------------------------------------------------------------------
    # 1) Build mapping run_ID → test_type from Tests_log.jsonl
    # ------------------------------------------------------------------
    run_to_type: Dict[str, str] = {}
    try:
        with open(tests_log_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                entry = json.loads(line)
                rid = (entry.get("run_ID") or entry.get("_runID"))
                ttype = entry.get("test_type")
                if rid and ttype:
                    run_to_type[str(rid)] = ttype
    except Exception as e:
        dbg("[submission] warn: cannot read tests_log:", e)

    # ------------------------------------------------------------------
    # 2) Find the run folder + DataEntry entry for this sample
    # ------------------------------------------------------------------
    found_run_folder: Optional[Path] = None
    data_entry: Optional[dict] = None
    for run_folder in runs_base.iterdir():
        if not run_folder.is_dir():
            continue
        de = run_folder / "DataEntry.json"
        if not de.exists():
            continue
        try:
            raw = json.loads(de.read_text(encoding="utf-8"))
        except Exception:
            continue
        entries = raw if isinstance(raw, list) else [raw]
        for e in entries:
            if (e.get("Sample ID") or e.get("sample_id")) == sample_id:
                found_run_folder = run_folder
                data_entry = e
                break
        if found_run_folder:
            break

    if not (found_run_folder and data_entry):
        dbg("[submission] no DataEntry match for sample:", sample_id)
        return

    rid = (data_entry.get("_runID") or data_entry.get("run_ID"))
    ttype = run_to_type.get(str(rid))
    if rid and ttype:
        canon = runs_base / f"{ttype}_{rid}"
        if canon.exists():
            found_run_folder = canon

    submission_id = (data_entry.get("Submission") or data_entry.get("submission"))
    if not submission_id:
        dbg("[submission] no Submission ID")
        return

    # ------------------------------------------------------------------
    # 3) DB-backed lookups via object_sql_db
    # ------------------------------------------------------------------
    project_path = nouns_base.parent.resolve()  # .../projects/<project>

    # 3a. Submission.received_date
    received_date = _load_submission_received_date(project_path, str(submission_id)) or ""
    if received_date:
        replace_placeholders_in_docx(doc, [("{{sampleReceived}}", received_date)])
    else:
        dbg("[submission] no received_date available from DB")

    # 3b. Picture adjectives (Sample / Potency Sample / Terpene Sample / Submission)
    picture_raw: Optional[str] = None

    # Try test-type-specific sample nouns first
    candidate_nouns: List[str] = []
    if ttype == "Potency_Test":
        candidate_nouns.append("Potency Sample")
    elif ttype == "Terpene_Test":
        candidate_nouns.append("Terpene Sample")

    # Generic Sample noun as additional fallback
    candidate_nouns.append("Sample")

    for noun_name in candidate_nouns:
        val = _load_noun_picture_value(project_path, noun_name, str(sample_id))
        if val:
            dbg(f"[submission] using picture from noun '{noun_name}'")
            picture_raw = val
            break

    # If none found on sample nouns, fall back to Submission's Picture adjective
    if not picture_raw:
        picture_raw = _load_noun_picture_value(project_path, "Submission", str(submission_id))
        if picture_raw:
            dbg("[submission] using picture from Submission noun")

    if not picture_raw:
        dbg("[submission] no picture found in DB; skipping sampleImage insertion")
        return

    # ------------------------------------------------------------------
    # 4) Resolve picture path relative to nouns_base and insert
    # ------------------------------------------------------------------
    s = str(picture_raw).strip()
    if s.startswith("nouns/"):
        image_rel = s.split("/", 1)[1]
    else:
        image_rel = s

    image_full_path = (nouns_base / image_rel).resolve()
    if not image_full_path.exists():
        dbg("[submission] image path does not exist on disk:", image_full_path)
        pic_path = None
    else:
        pic_path = image_full_path

    for p in doc.paragraphs:
        if replace_placeholder_across_runs(p.runs, "{{sampleImage}}", pic_path, 50, 50):
            dbg("[submission] inserted image in a paragraph")
    for tbl in doc.tables:
        process_table(tbl, "{{sampleImage}}", pic_path, 50, 50)

    dbg("[submission] done")

def compute_two_step_placeholders(weights_data: dict, coa_name_map_path: Path) -> list[tuple[str, str]]:
    """
    Build placeholder pairs for the 'two-step' fields:
      {{amountTHCA}}, {{amountD9}}, {{amountCBD}}, {{amountTotalCan}}, {{amountTotalTerp}}

    Uses the COA Name Map to resolve display names -> instrument codes, then pulls
    numbers from weights_data['potency'] / weights_data['terpene'].
    """
    # Default empty if missing inputs
    if not (weights_data and coa_name_map_path and coa_name_map_path.exists()):
        return []

    # Load the name map (display names → instrument codes)
    with coa_name_map_path.open("r", encoding="utf-8") as f:
        name_map = [json.loads(line) for line in f if line.strip()]

    # Build lookup of display name (lowercased) → instrument code for cannabinoids
    display_to_code_can = {}
    display_to_code_terp = {}
    for entry in name_map:
        disp = (entry.get("COA Display Name") or "").strip()
        code = (entry.get("Instrument Code") or "").strip()
        atype = (entry.get("Analyte Type") or "").strip()
        if not disp or not code:
            continue
        if atype == "Cannabinoid":
            display_to_code_can[disp.lower()] = code
        elif atype == "Terpene":
            display_to_code_terp[disp.lower()] = code

    def _find_code(substr: str, pool: dict[str, str]) -> str | None:
        """Find first instrument code whose display name contains the substring (case-insensitive)."""
        s = substr.lower()
        for disp_lc, code in pool.items():
            if s in disp_lc:
                return code
        return None

    potency = weights_data.get("potency", {}) or {}
    terpene = weights_data.get("terpene", {}) or {}

    # Totals come from build_weights_for_sample() as '<key>_total'
    total_can_raw  = potency.get("potency_total")
    total_terp_raw = terpene.get("terpene_total")

    # Individual targets
    thca_code = _find_code("thca", display_to_code_can)
    d9_code   = _find_code("delta-9", display_to_code_can) or _find_code("delta-9", display_to_code_can) or _find_code("delta 9", display_to_code_can)
    cbd_code  = _find_code("cbd)", display_to_code_can) or _find_code("cannabidiol", display_to_code_can)

    thca_val = potency.get(thca_code) if thca_code else None
    d9_val   = potency.get(d9_code)   if d9_code   else None
    cbd_val  = potency.get(cbd_code)  if cbd_code  else None

    def _fmt_num(x):
        try:
            v = float(x)
            if v == 0:
                return "ND"
            return f"{v:.2f}"
        except (TypeError, ValueError):
            return "NT"

    pairs = [
        ("{{amountTHCA}}",      _fmt_num(thca_val)),
        ("{{amountD9}}",        _fmt_num(d9_val)),
        ("{{amountCBD}}",       _fmt_num(cbd_val)),
        ("{{amountTotalCan}}",  _fmt_num(total_can_raw)),
        ("{{amountTotalTerp}}", _fmt_num(total_terp_raw)),
    ]
    return pairs
