# docker/coa_generator/coa_generator.py

import sys
from pathlib import Path
import json
import shutil
from datetime import datetime
from collections import defaultdict
from docx import Document
import csv
from collections import defaultdict
from docx.shared import Pt, Mm

# Add project root to sys.path
project_root = Path(__file__).parent.parent.resolve()
sys.path.insert(0, str(project_root))

from utils.interface import menu_prompt, indexed_choice
from utils.status import get_status_breakdown

def run_pre_pphrase() -> dict:
    """
    Pre-run for COA generator.
    Uses collect_completed_samples to find completed Potency_Test and Terpene_Test runs,
    displays indexed options with client info, and allows interactive sample selection.
    """

    import os

    cwd = Path.cwd()
    print(f"🔍 Current working directory: {cwd}")

    # 🔧 Flip this to enable/disable directory debug mode
    directory_mode = False

    if directory_mode:
        print("📂 Full directory tree:")
        for root, dirs, files in os.walk(".", topdown=True):
            level = root.replace(os.getcwd(), '').count(os.sep)
            indent = ' ' * 4 * level
            print(f"{indent}{root}/")
            subindent = ' ' * 4 * (level + 1)
            for f in files:
                print(f"{subindent}{f}")

        # ❌ TEMPORARY EXIT to inspect mount correctness
        raise RuntimeError("💥 DEBUG STOP – inspect directory tree above before proceeding.")

    # 1) Collect completed samples
    completed_samples = collect_completed_samples(cwd)

    # 2) Filter only Potency_Test and Terpene_Test
    filtered = [
        {
            "sample_id": s["sample_id"],
            "run_id": s["run_id"],
            "verb": s["test_type"],
            "client": s["client"],
            "sample_name": s.get("sample_name"),
            "sample_type": s.get("sample_type")
        }
        for s in completed_samples
        if s["test_type"] in ("Potency_Test", "Terpene_Test")
    ]

    # 3) De-dupe
    seen = set()
    unique_entries: list[dict[str, str]] = []
    for entry in filtered:
        key = (entry["sample_id"], entry["verb"], entry["run_id"])
        if key not in seen:
            seen.add(key)
            unique_entries.append(entry)

    # 4) Interactive pick/remove
    catalog = unique_entries.copy()
    selected: list[dict[str, str]] = []

    while True:
        print("\n📋 Available samples:")
        if not catalog:
            print(" (no more available samples)")
        else:
            for i, e in enumerate(catalog):
                print(f"[{i}] {e['sample_id']} ({e['verb']}: {e['run_id']}) – Client: {e['client']}")

        print("\n📋 Selected samples:")
        if not selected:
            print(" (none)")
        else:
            for s in selected:
                print(f" - {s['sample_id']} ({s['verb']}:{s['run_id']}) – Client: {s['client']}")

        action = menu_prompt({"a": "Add sample", "d": "Delete sample", "q": "Quit selection"})
        if action == "a":
            if not catalog:
                print("⚠️ No options available to add.")
                continue
            idx = indexed_choice(
                [f"{e['sample_id']} ({e['verb']}: {e['run_id']}) – Client: {e['client']}" for e in catalog],
                "Select a sample to add:"
            )
            if idx is not None:
                selected.append(catalog.pop(idx))

        elif action == "d":
            if not selected:
                print("⚠️ No samples to delete.")
                continue
            idx = indexed_choice(
                [f"{e['sample_id']} ({e['verb']}: {e['run_id']}) – Client: {e['client']}" for e in selected],
                "Select a sample to delete:"
            )
            if idx is not None:
                catalog.append(selected.pop(idx))

        else:  # 'q'
            break

    # 5) Return the selected samples only
    return {"samples": selected}
    
def run_pphrase(output_dir, payload=None):
    print("👋 Hello world from prepositional phrase runner!")
    print(f"📤 Output directory: {output_dir}")

    if not payload or not payload.get("samples"):
        raise ValueError("❌ No payload provided or payload is empty – cannot proceed with COA generation.")

    print(f"📦 Received payload: {payload}")

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # 🔥 Read Primary Aromas entries at start
    coa_name_map_dir = Path("/app/inputs/Primary Aromas")  # Adjust for container path
    coa_file = coa_name_map_dir / "items.jsonl"

    if not coa_file.exists():
        print(f"❌ Primary Aromas file not found at {coa_file}")
        coa_entries = []
    else:
        coa_entries = []
        with coa_file.open("r") as f:
            for line in f:
                line = line.strip()
                if line:
                    entry = json.loads(line)
                    coa_entries.append(entry)

    print("\n📦 Primary Aromas entries loaded:")
    for entry in coa_entries:
        print(entry)

    # Load run_ID → date_tested mapping
    tests_log_path = Path("/app/types/Tests_log.jsonl")
    runid_dates = load_runid_dates(tests_log_path)

    # Load run_ID → test_type mapping
    tests_log_path = Path("/app/types/Tests_log.jsonl")
    run_to_type = {}
    if tests_log_path.exists():
        with tests_log_path.open() as f:
            for line in f:
                entry = json.loads(line)
                rid = entry.get("run_ID")
                ttype = entry.get("test_type")
                if rid and ttype:
                    run_to_type[rid] = ttype
    else:
        print(f"❌ Tests_log.jsonl not found at {tests_log_path}")

    inputs_base = Path("/app/inputs")  # Adjust this if your mounted inputs differ

    # 🔥 Consolidate payload samples into combined sample records
    samples_consolidated = consolidate_sample_tests(payload, runid_dates)
    print(f"🔎 Consolidated samples: {samples_consolidated}")

    # Get planned output structure
    output_structure = get_output_structure(payload)

    # Create folders
    for folder in output_structure.get("folders", []):
        path = output_dir / folder
        path.mkdir(parents=True, exist_ok=True)
        print(f"📁 Created folder: {path}")

    # Copy and rename COA template per sample
    template_src = Path("/app/docker_inputs/Python COA Template.docx")

    if not template_src.exists():
        print(f"❌ Template not found at {template_src}")
        return

    for sample in samples_consolidated:
        sample_id = sample["sample_id"]
        client = sample.get("client", "UnknownClient")
        date = datetime.today().strftime("%Y%m%d")

        folder_path = output_dir / f"{client}/{date}"
        file_name = f"COA for {sample_id}.docx"
        dst = folder_path / file_name

        shutil.copy(template_src, dst)
        print(f"✅ Copied template to {dst}")

        # Open document
        doc = Document(dst)

        # 🔥 Get weights data for this consolidated sample
        weights_data = print_weights_for_sample(sample, inputs_base)

        # 🔥 Call analyte mapping subroutine
        analyte_table_data = build_analyte_table_data(
            weights_data,
            inputs_base / "COA Name Map/items.jsonl",
            Path("docker_inputs/Python COA Template.docx")
        )

        # 🔥 Insert analyte data into docx tables
        insert_analyte_table_data(doc, analyte_table_data)

        # 🔥 Insert primary aromas into COA template
        # 🔥 Load COA Name Map entries 
        coa_name_map_file = Path("/app/inputs/COA Name Map/items.jsonl")
        with coa_name_map_file.open("r") as f:
            coa_name_map = [json.loads(line.strip()) for line in f if line.strip()]
        primary_aroma_file = Path("/app/inputs/Primary Aromas/items.jsonl")
        #🔥 and primary aromas before calling
        primary_aroma_pictures = []
        with primary_aroma_file.open("r") as f:
            for line in f:
                if line.strip():
                    primary_aroma_pictures.append(json.loads(line))

        # 🔥 Call insertion
        insert_primary_aromas(doc, weights_data, coa_name_map, primary_aroma_pictures, inputs_base)

        # 🔥 Insert sample submission data into COA
        submissions_path = inputs_base / "Submission"
        insert_sample_submission_data(
            doc=doc,
            sample_id=sample_id,
            submissions_path=submissions_path,
            inputs_base=inputs_base,           # optional, defaults to /app/inputs
            tests_log_path=Path("/app/types/Tests_log.jsonl")  # optional override
        )

        # 🔥 Build replacements list
        terpene_info = sample.get("terpene", {})
        potency_info = sample.get("potency", {})

        terpene_date_tested = terpene_info.get("date_tested", "n/a")
        potency_date_tested = potency_info.get("date_tested", "n/a")
        sample_name = sample.get("sample_name", "n/a")
        sample_type = sample.get("sample_type", "n/a")

        # Extract from weights_data
        terpene_total = weights_data.get("terpene", {}).get("terpene_total", "NT")
        potency_total = weights_data.get("potency", {}).get("potency_total", "NT")
        amount_thca = weights_data.get("potency", {}).get("THCA", "NT")
        amount_d9 = weights_data.get("potency", {}).get("d9", "NT")
        amount_cbd = weights_data.get("potency", {}).get("CBD", "NT")

        # Convert numeric values to floats with 2 decimal places, else keep as NT/ND
        def safe_format(val):
            try:
                num = float(val)
                if num == 0:
                    return "ND"
                return f"{num:.2f}"
            except (ValueError, TypeError):
                return val

        # Build final replacements list
        replacements = [
            ("{{Client}}", client),
            ("{{sampleID}}", sample_id),
            ("{{sampleName}}", sample_name),
            ("{{sampleType}}", sample_type),
            ("{{terpeneDateTested}}", terpene_date_tested),
            ("{{cannabinoidDateTested}}", potency_date_tested),
            ("{{amountTotalTerp}}", safe_format(terpene_total)),
            ("{{amountTotalCan}}", safe_format(potency_total)),
            ("{{amountTHCA}}", safe_format(amount_thca)),
            ("{{amountD9}}", safe_format(amount_d9)),
            ("{{amountCBD}}", safe_format(amount_cbd)),
        ]

        # 🔧 Replace placeholders using subroutine
        replace_placeholders_in_docx(doc, replacements)

        # Save updated document
        doc.save(dst)
        print(f"✏️ Replaced placeholders and saved {dst}")

def get_metadata():
    return {
        "name": "coa_generator",
        "entrypoint": "coa_generator.py",
        "dependencies": ["python-docx"]
    }

def get_io_manifest():
    return {
        "Potency_Test": {"type": "verb"},
        "Terpene_Test": {"type": "verb"},
        "COA Name Map": {"type": "noun"},
        "Primary Aromas": {"type": "noun"},
        "Submission": {"type": "noun"},
    }

def get_output_structure(context):
    """
    Defines the output structure per client without runID subfolders.

    Returns:
        dict with folders and files to create.
    """

    samples = context.get("samples", [])

    # Group samples by client
    client_samples = defaultdict(list)
    for s in samples:
        client = s.get("client", "UnknownClient")
        client_samples[client].append(s)

    date = context.get("date")
    if not date:
        date = datetime.today().strftime("%Y%m%d")

    folders = []
    files = {}

    for client, sample_list in client_samples.items():
        client_folder = f"{client}/{date}"
        folders.append(client_folder)

        # Map files to create for each sample in the client/date folder
        for sample in sample_list:
            path = client_folder
            files[path] = files.get(path, []) + [f"{sample['sample_id']}_COA.pdf"]

    return {
        "folders": folders,
        "files": files
    }

def define_outputs(context):
    """
    Defines the output files this prepositional phrase will create.

    Returns:
        list of dicts, each representing a file with:
            - name: filename only
            - type: file type (e.g. pdf, csv, json)
            - metadata: optional extra data for generation
    """

    runs = context.get("runs", {})
    outputs = []

    for run_id, samples in runs.items():
        for sample_id in samples:
            outputs.append({
                "name": f"{sample_id}_COA.pdf",
                "type": "pdf",
                "metadata": {
                    "sample_id": sample_id,
                    "run_id": run_id
                }
            })

    return outputs

def is_run_complete(status_breakdown: dict) -> bool:
    """
    Determines if a run is considered complete for COA generation.
    Requires:
        - raw_data == "Uploaded"
        - data_entry == "Complete"
        - interpretation == "Parsed"
        - adverb_info == "Complete"
    """
    return (
        status_breakdown.get("raw_data") == "Uploaded"
        and status_breakdown.get("data_entry") == "Complete"
        and status_breakdown.get("interpretation") == "Parsed"
        and status_breakdown.get("adverb_info") == "Complete"
    )

def manage_list_prompt(item_label: str) -> list:
    """
    Interactive menu prompt to build a list by adding or deleting items.

    Args:
        item_label: descriptive label for the items (e.g. "sample", "client")

    Returns:
        list of items entered by the user
    """
    items = []

    while True:
        print("\n📋 Current list:")
        if items:
            for i, it in enumerate(items, 1):
                print(f" {i}. {it}")
        else:
            print(" (empty)")

        choice = menu_prompt({
            "a": f"Add {item_label}",
            "d": f"Delete {item_label}",
            "q": "Quit"
        })

        if choice == "a":
            new_item = input(f"Enter {item_label}: ").strip()
            if new_item:
                items.append(new_item)
                print(f"✅ Added {item_label}: {new_item}")
            else:
                print("❌ No input entered.")

        elif choice == "d":
            if not items:
                print("❌ List is empty. Nothing to delete.")
                continue

            try:
                idx = int(input(f"Enter number to delete (1-{len(items)}): ").strip())
                if 1 <= idx <= len(items):
                    removed = items.pop(idx - 1)
                    print(f"🗑️ Removed {item_label}: {removed}")
                else:
                    print("❌ Invalid number.")
            except ValueError:
                print("❌ Invalid input.")

        elif choice == "q":
            break

    return items

def load_full_verb_def(project_path: Path, verb_key: str) -> dict:
    """
    Loads the entire verb definition (not just data_entry_schema)
    from verb_types.json.
    """
    with open(project_path / "verb_types.json") as vf:
        all_verbs = json.load(vf)
    return all_verbs.get(verb_key, {})

def collect_completed_samples(project_root: Path) -> list[dict]:
    """
    Scans the inputs/ directory for completed runs,
    extracts sample IDs, run IDs, and client from DataEntry.json,
    and attaches the corresponding test_type from Tests_log.jsonl.
    """

    inputs_dir = project_root / "inputs"
    tests_log_path = project_root / "types" / "Tests_log.jsonl"

    # Load Tests_log.jsonl into a lookup dict: run_ID -> test_type
    runid_to_type = {}
    with open(tests_log_path) as f:
        for line in f:
            entry = json.loads(line)
            runid_to_type[entry["run_ID"]] = entry["test_type"]

    # Define the exact breakdown dict for comparison
    required_breakdown = {
        "raw_data": "Uploaded",
        "data_entry": "Complete",
        "interpretation": "Parsed",
        "adverb_info": "Complete"
    }

    results = []

    # Iterate folders in inputs/
    for folder in inputs_dir.iterdir():
        if not folder.is_dir():
            continue

        status_path = folder / "Status.json"
        dataentry_path = folder / "DataEntry.json"

        # Check if Status.json exists and has the required breakdown
        if status_path.exists():
            with open(status_path) as f:
                status = json.load(f)
            if status.get("breakdown") == required_breakdown:

                # If DataEntry.json exists, load and process
                if dataentry_path.exists():
                    with open(dataentry_path) as f:
                        entries = json.load(f)

                    for entry in entries:
                        run_id = entry.get("_runID")
                        sample_id = entry.get("Sample ID")
                        client = entry.get("Client")
                        sample_name = entry.get("Sample Name")
                        sample_type = entry.get("Sample Type")

                        # Get test_type from Tests_log.jsonl lookup
                        test_type = runid_to_type.get(run_id)

                        results.append({
                            "sample_id": sample_id,
                            "run_id": run_id,
                            "test_type": test_type,
                            "client": client,
                            "sample_name": sample_name,
                            "sample_type": sample_type
                        })

    return results

def print_weights_for_sample(sample_or_list, inputs_base):
    """
    For a single consolidated sample dict or list of them:
    - Loads weights data per test type
    - Combines into a nested dict
    - Calculates totals per type
    - Prints the combined JSON-like output
    """

    # If input is a list, process each sample recursively
    if isinstance(sample_or_list, list):
        for sample in sample_or_list:
            print_weights_for_sample(sample, inputs_base)
        return

    sample = sample_or_list
    sample_id = sample["sample_id"]

    combined_data = {"sample": sample_id}

    for test_type in ["terpene", "potency"]:
        if test_type in sample:
            run_id = sample[test_type]["run_id"]
            verb = "Terpene_Test" if test_type == "terpene" else "Potency_Test"

            folder_name = f"{verb}_{run_id}"
            run_folder = inputs_base / folder_name
            weights_csv_path = run_folder / "Weights.csv"

            if not weights_csv_path.exists():
                print(f"❌ Weights.csv not found for run {folder_name}")
                continue

            with weights_csv_path.open("r", newline='') as f:
                reader = csv.reader(f)
                rows = list(reader)

            if not rows:
                print(f"⚠️ No data in Weights.csv for run {folder_name}")
                continue

            headers = rows[0]

            sample_row = None
            for row in rows[1:]:
                if row[0].strip() == sample_id:
                    sample_row = row
                    break

            if sample_row:
                # Build a dict mapping header to value (excluding Sample column)
                data_dict = dict(zip(headers[1:], sample_row[1:]))

                # Calculate total (sum of all numeric analyte values)
                total = 0
                for val in data_dict.values():
                    try:
                        num = float(val)
                        total += num
                    except (ValueError, TypeError):
                        continue

                # Add total to data_dict
                data_dict[f"{test_type}_total"] = round(total, 3)
                combined_data[test_type] = data_dict
            else:
                print(f"⚠️ Sample ID {sample_id} not found in Weights.csv for run {folder_name}")

    # Print the final combined structured data
    print("\n📦 Combined JSON data:")
    print(combined_data)

    return combined_data

def consolidate_sample_tests(payload, runid_dates):
    """
    Groups payload entries by sample_id, ensures only one entry per test_type,
    and returns consolidated dicts for each sample.
    """

    # Group entries by sample_id
    grouped = defaultdict(list)
    for s in payload["samples"]:
        grouped[s["sample_id"]].append(s)

    consolidated = []

    for sample_id, entries in grouped.items():
        # Further group by test_type within this sample_id
        test_type_map = defaultdict(list)
        for entry in entries:
            test_type_map[entry["verb"]].append(entry)

        sample_record = {"sample_id": sample_id}

        # 🛠 Extract client, sample_name, sample_type from any entry (assuming same for this sample_id)
        client = entries[0].get("client", "UnknownClient")
        sample_name = entries[0].get("sample_name", "UnknownName")
        sample_type = entries[0].get("sample_type", "UnknownType")

        sample_record["client"] = client
        sample_record["sample_name"] = sample_name
        sample_record["sample_type"] = sample_type

        for test_type, test_entries in test_type_map.items():
            if len(test_entries) == 1:
                chosen = test_entries[0]
            else:
                # ⚠️ Multiple entries for same test_type – resolve interactively
                print(f"\n⚠️ Multiple entries found for sample {sample_id} test {test_type}. Please resolve:")
                while len(test_entries) > 1:
                    for idx, e in enumerate(test_entries):
                        print(f"[{idx}] run_id: {e['run_id']}, client: {e.get('client','Unknown')}")

                    choice = input(f"Select which entry to KEEP for sample {sample_id} test {test_type} (or 'd' to delete one): ").strip()

                    if choice.lower() == 'd':
                        del_idx = input("Enter index to DELETE: ").strip()
                        if del_idx.isdigit() and int(del_idx) < len(test_entries):
                            removed = test_entries.pop(int(del_idx))
                            print(f"❌ Removed entry with run_id {removed['run_id']}")
                        else:
                            print("⚠️ Invalid index.")
                    elif choice.isdigit() and int(choice) < len(test_entries):
                        chosen = test_entries[int(choice)]
                        break
                    else:
                        print("⚠️ Invalid input.")

                if len(test_entries) == 1:
                    chosen = test_entries[0]
                else:
                    raise ValueError(f"❌ Resolution failed for sample {sample_id} test {test_type}")

            # Map chosen entry under test_type key
            run_id = chosen["run_id"]
            date_tested = runid_dates.get(run_id)

            if test_type == "Potency_Test":
                sample_record["potency"] = {
                    "run_id": run_id,
                    "date_tested": date_tested
                }
            elif test_type == "Terpene_Test":
                sample_record["terpene"] = {
                    "run_id": run_id,
                    "date_tested": date_tested
                }
            else:
                print(f"⚠️ Unknown test_type {test_type} for sample {sample_id}")

        consolidated.append(sample_record)

    return consolidated

def replace_placeholders_in_docx(doc: Document, replacements: list[tuple[str, str]]):
    """
    Replaces placeholders in a docx Document, including nested tables.
    """

    def replace_in_paragraph(paragraph, replacements):
        for placeholder, replacement in replacements:
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)

    def replace_in_table(table, replacements):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
                for nested_table in cell.tables:
                    replace_in_table(nested_table, replacements)

    # Replace in document paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    # Replace in all tables (including nested)
    for table in doc.tables:
        replace_in_table(table, replacements)

def load_runid_dates(tests_log_path):
    runid_dates = {}
    with open(tests_log_path) as f:
        for line in f:
            entry = json.loads(line)
            runid_dates[entry["run_ID"]] = entry.get("date_tested")
    return runid_dates

def extract_analytes_by_table(doc_path: Path, display_to_entry: dict):
    """
    Recursively scan all tables (including nested) and return a list of rows:
    { table_index, analyte, instrument_code, analyte_type }
    Includes rows where analyte == "Total".
    """
    doc = Document(doc_path)
    rows = []

    def parse_table(table, index):
        for tr in table.rows:
            name = tr.cells[0].text.strip()
            if name and (name in display_to_entry or name.lower() == "total"):
                entry = display_to_entry.get(name, {})
                rows.append({
                    "table_index": index,
                    "analyte": name,
                    "instrument_code": entry.get("Instrument Code"),
                    "analyte_type": entry.get("Analyte Type")
                })
            # Recurse into any nested tables in this row’s cells
            for cell in tr.cells:
                for nested in cell.tables:
                    parse_table(nested, f"{index}-nested")

    # Kick off with every top-level table
    for i, tbl in enumerate(doc.tables):
        parse_table(tbl, str(i))

    return rows

def build_analyte_table_data(weights_data: dict,
                             coa_name_map_path: Path,
                             coa_template_path: Path):
    """
    Returns list of dict rows ready for insertion, including totals.
    """
    # 1. Load name map
    with open(coa_name_map_path) as f:
        name_map = [json.loads(line) for line in f]
    display_to_entry = {e["COA Display Name"]: e for e in name_map}

    # 2. Extract analyte + Total rows
    extracted = extract_analytes_by_table(coa_template_path, display_to_entry)

    # 3. Map analyte types to weights_data keys
    type_to_key = {"Terpene": "terpene", "Cannabinoid": "potency"}

    last_type = None
    for row in extracted:
        analyte = row["analyte"]

        if analyte.lower() == "total":
            # Total row → insert the precomputed total for last_type
            if last_type:
                key = type_to_key[last_type]
                total_key = f"{key}_total"
                raw_tot = weights_data.get(key, {}).get(total_key)
                if raw_tot is None:
                    pct, mg = "NT", "NT"
                else:
                    val = float(raw_tot)
                    pct = val
                    mg  = round(val * 10, 3)
            else:
                pct, mg = "NT", "NT"

        else:
            # Regular analyte → compute from weights_data
            atype = row["analyte_type"]
            code  = row["instrument_code"]
            last_type = atype  # remember for next “Total”

            key = type_to_key.get(atype)
            raw = weights_data.get(key, {}).get(code)

            if raw is None or raw == "NaN":
                pct, mg = "NT", "NT"
            else:
                val = float(raw)
                if val == 0:
                    pct, mg = "ND", "ND"
                else:
                    pct = val
                    mg  = round(val * 10, 3)

        row["%"]   = pct
        row["mg/g"] = mg

    # 4. Debug print
    print("\n=== BUILT ANALYTE TABLE DATA ===")
    for r in extracted:
        print(r)

    return extracted

def insert_analyte_table_data(doc: Document, analyte_table_data: list[dict]):
    """
    Inserts analyte % and mg/g values into their respective docx table rows.
    Sets font size to 8 pt for inserted values.
    """
    def set_font_size(cell, size_pt):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(size_pt)

    for row_data in analyte_table_data:
        table_index = row_data["table_index"]
        analyte = row_data["analyte"]
        percent = row_data["%"]
        mg_per_g = row_data["mg/g"]

        # If nested table, parse index like "1-nested"
        if isinstance(table_index, str) and "-nested" in table_index:
            main_index, _ = table_index.split("-nested")
            table = doc.tables[int(main_index)]

            for row in table.rows:
                for cell in row.cells:
                    for nested_table in cell.tables:
                        for t_row in nested_table.rows:
                            if t_row.cells[0].text.strip() == analyte:
                                if len(t_row.cells) >= 3:
                                    t_row.cells[1].text = str(percent)
                                    t_row.cells[2].text = str(mg_per_g)
                                    set_font_size(t_row.cells[1], 8)
                                    set_font_size(t_row.cells[2], 8)
                                elif len(t_row.cells) == 2:
                                    t_row.cells[1].text = str(percent)
                                    set_font_size(t_row.cells[1], 8)
                                print(f"✅ Filled nested {analyte}: %={percent}, mg/g={mg_per_g}")
        else:
            # Top-level table
            table = doc.tables[int(table_index)]
            for t_row in table.rows:
                if t_row.cells[0].text.strip() == analyte:
                    if len(t_row.cells) >= 3:
                        t_row.cells[1].text = str(percent)
                        t_row.cells[2].text = str(mg_per_g)
                        set_font_size(t_row.cells[1], 8)
                        set_font_size(t_row.cells[2], 8)
                    elif len(t_row.cells) == 2:
                        t_row.cells[1].text = str(percent)
                        set_font_size(t_row.cells[1], 8)
                    print(f"✅ Filled {analyte}: %={percent}, mg/g={mg_per_g}")

def insert_primary_aromas(
    doc: Document,
    weights_data: dict,
    coa_name_map: list[dict],
    primary_aroma_pictures: list[dict],
    inputs_base: Path = Path("inputs"),
):
    """
    Inserts top unique primary aromas based on highest terpene concentrations,
    replacing both aroma and picture placeholders in-place (now limited to top 3).
    """

    # 1) Build instrument code → aroma mapping from COA Name Map
    instrument_to_aroma = {
        entry["Instrument Code"]: entry["Primary Aroma"]
        for entry in coa_name_map
        if entry.get("Analyte Type") == "Terpene" and entry.get("Primary Aroma")
    }

    # 2) Build aroma → picture path mapping
    aroma_to_path = {}
    for entry in primary_aroma_pictures:
        aroma = entry.get("Aroma")
        raw_id = entry.get("Picture ID")
        if aroma and raw_id:
            rel_path = raw_id.split("/", 1)[1] if raw_id.startswith("nouns/") else raw_id
            full_path = inputs_base / rel_path
            aroma_to_path[aroma] = full_path

    # 3) Get top 3 terpene codes by concentration
    terpene_data = weights_data.get("terpene", {})
    terpene_concs = []
    for code, val in terpene_data.items():
        if code in instrument_to_aroma and val not in ("NaN", "NT"):
            try:
                terpene_concs.append((code, float(val)))
            except ValueError:
                pass

    # Sort descending by concentration and build unique aromas
    top3 = sorted(terpene_concs, key=lambda x: x[1], reverse=True)[:3]
    unique_aromas = []
    for code, _ in top3:
        aroma = instrument_to_aroma.get(code)
        if aroma and aroma not in unique_aromas:
            unique_aromas.append(aroma)
        if len(unique_aromas) >= 3:
            break

    # 4) Replace primaryAroma# placeholders using centralized function
    aroma_replacements = []
    for i in range(1, 4):
        aroma_ph = f"{{{{primaryAroma{i}}}}}"
        aroma_text = unique_aromas[i-1] if i <= len(unique_aromas) else ""
        aroma_replacements.append((aroma_ph, aroma_text))

    replace_placeholders_in_docx(doc, aroma_replacements)

    # 5) Replace primaryPicture# placeholders with images in-place
    for i in range(1, 4):
        picture_ph = f"{{{{primaryPicture{i}}}}}"
        aroma_text = unique_aromas[i-1] if i <= len(unique_aromas) else ""
        pic_path = aroma_to_path.get(aroma_text)

        # Scan all paragraphs in the main document body
        for p in doc.paragraphs:
            replace_placeholder_across_runs(p.runs, picture_ph, pic_path, 20, 20)

        # Scan all tables recursively using process_table
        for tbl in doc.tables:
            process_table(tbl, picture_ph, pic_path, 20, 20)

def replace_placeholder_across_runs(runs, placeholder, pic_path, widthSize=None, heightSize=None):
    """
    Searches for placeholder across runs.
    Deletes it if found and inserts image at the first run containing any part.
    Returns True if replaced, False otherwise.
    """
    full_text = "".join(run.text for run in runs)
    if placeholder not in full_text:
        return False

    # Track how much of the placeholder remains to delete
    remaining = placeholder
    for run in runs:
        if not remaining:
            break
        if remaining in run.text:
            run.text = run.text.replace(remaining, "")
            remaining = ""
        elif remaining.startswith(run.text):
            remaining = remaining[len(run.text):]
            run.text = ""
        else:
            pass

    # Insert image at first run if file exists
    if pic_path and pic_path.exists():
        # Default size logic for each dimension independently
        if widthSize is None:
            widthSize = 50
        if heightSize is None:
            heightSize = 50

        runs[0].add_picture(
            str(pic_path),
            width=Mm(widthSize),
            height=Mm(heightSize)
        )

    return True

def process_cell(cell, picture_ph, pic_path, widthSize, heightSize):
    # Process paragraphs in this cell
    for p in cell.paragraphs:
        replace_placeholder_across_runs(p.runs, picture_ph, pic_path, widthSize, heightSize)

    # Recurse into nested tables within this cell
    for tbl in cell.tables:
        process_table(tbl, picture_ph, pic_path, widthSize, heightSize)


def process_table(tbl, picture_ph, pic_path, widthSize, heightSize):
    for row in tbl.rows:
        for cell in row.cells:
            process_cell(cell, picture_ph, pic_path, widthSize, heightSize)

def insert_sample_submission_data(
    doc: Document,
    sample_id: str,
    submissions_path: Path,
    inputs_base: Path = Path("/app/inputs"),
    tests_log_path: Path = Path("/app/types/Tests_log.jsonl"),
):
    """
    Inserts sample received date and submission image into COA document.

    :param doc: The COA docx Document object.
    :param sample_id: Sample ID to find in data dump.
    :param submissions_path: Path to Submission folder (inputs/Submission).
    :param inputs_base: Base inputs folder containing run folders.
    :param tests_log_path: Path to Tests_log.jsonl for run→test_type lookup.
    """
    print("\n🔧 Running insert_sample_submission_data...")
    print(f"🔎 Sample ID: {sample_id}")
    print(f"📁 Submissions path: {submissions_path}")
    print(f"📁 Inputs base: {inputs_base}")
    print(f"📁 Tests_log.jsonl: {tests_log_path}")

    # ─── 1) Load the DataEntry.json inside the correct run folder ───
    # First, scan all mounted run folders for the sample’s DataEntry.json
    # We know each run folder is named "<test_type>_<run_ID>"
    # But we need run_ID and test_type, so scan Tests_log.jsonl

    # Load Tests_log.jsonl into dict: run_ID → test_type
    run_to_type = {}
    if not tests_log_path.exists():
        print(f"❌ Tests_log.jsonl not found at {tests_log_path}")
        return
    with open(tests_log_path) as f:
        for line in f:
            entry = json.loads(line)
            rid = entry.get("run_ID")
            ttype = entry.get("test_type")
            if rid and ttype:
                run_to_type[rid] = ttype

    # Now we don’t yet know run_ID until we inspect each folder;
    # Instead, glob all possible run folders, open each DataEntry.json,
    # and match on sample_id.

    found = False
    for run_folder in inputs_base.iterdir():
        if not run_folder.is_dir():
            continue

        data_entry_f = run_folder / "DataEntry.json"
        if not data_entry_f.exists():
            continue

        with data_entry_f.open() as f:
            raw = json.load(f)

        # unify into a list of entry-dicts
        entries = raw if isinstance(raw, list) else [raw]

        for entry in entries:
            if entry.get("Sample ID") == sample_id:
                found = True
                print(f"✅ Found DataEntry.json entry for sample {sample_id} in run folder: {run_folder.name}")
                data_entry = entry
                break

        if found:
            break

    if not found:
        print(f"❌ Could not locate DataEntry.json for sample {sample_id} under {inputs_base}")
        return

    # ─── 2) Extract run_ID and verify test_type ───
    run_id = data_entry.get("_runID")
    print(f"🔎 Extracted _runID from JSON: {run_id}")
    if not run_id:
        print("❌ _runID missing in DataEntry.json")
        return

    test_type = run_to_type.get(run_id)
    if not test_type:
        print(f"❌ run_ID '{run_id}' not found in Tests_log.jsonl")
        return

    # Reconstruct the canonical run_folder and reload DataEntry.json to confirm
    canon_run_folder = inputs_base / f"{test_type}_{run_id}"
    print(f"🔄 Canonical run folder path: {canon_run_folder}")
    data_entry_f = canon_run_folder / "DataEntry.json"
    if not data_entry_f.exists():
        print(f"❌ DataEntry.json not found at canonical path: {data_entry_f}")
        return
    with open(data_entry_f) as f:
        data_entry = json.load(f)
    print(f"✅ Re-loaded DataEntry.json contents: {data_entry}")

    # after computing canon_run_folder…
    data_entry_f = canon_run_folder / "DataEntry.json"
    with data_entry_f.open() as f:
        raw = json.load(f)
    entries = raw if isinstance(raw, list) else [raw]
    # pick the first matching one
    data_entry = next((e for e in entries if e.get("Sample ID") == sample_id), None)
    if data_entry is None:
        print(f"❌ No matching Sample ID {sample_id} in {data_entry_f}")
        return

    # ─── 3) Extract Submission ID ───
    submission_id = data_entry.get("Submission")
    print(f"🔎 Extracted Submission ID: {submission_id}")
    if not submission_id:
        print(f"❌ Submission ID missing for sample {sample_id}")
        return

    # ─── 4) Load submissions lookup ───
    items_f = submissions_path / "items.jsonl"
    print(f"📄 Loading submissions from: {items_f}")
    if not items_f.exists():
        print("❌ items.jsonl not found in Submission folder")
        return

    submission_lookup = {}
    with open(items_f) as f:
        for line in f:
            item = json.loads(line)
            sid = item.get("submission_id")
            if sid:
                submission_lookup[sid] = item

    print(f"✅ Loaded {len(submission_lookup)} submissions; IDs: {list(submission_lookup.keys())}")
    sub_item = submission_lookup.get(submission_id)
    if not sub_item:
        print(f"❌ Submission ID {submission_id} not in items.jsonl")
        return
    print(f"✅ Retrieved submission entry: {sub_item}")

    # ─── 5) Extract received_date & raw image path ───
    received_date   = sub_item.get("received_date", "")
    image_path_raw  = sub_item.get("image", "")
    print(f"📅 Received date: {received_date}")
    print(f"🖼️ Raw image path: {image_path_raw}")

    # ─── 6) Resolve image path ───
    if image_path_raw.startswith("nouns/"):
        image_rel = image_path_raw.replace("nouns/", "", 1)
    else:
        image_rel = image_path_raw
    image_full_path = inputs_base / image_rel
    print(f"📁 Resolved image to: {image_full_path}")
    if not image_full_path.exists():
        print(f"⚠️ Image not found at resolved path: {image_full_path}")

    # ─── 7) Replace {{sampleReceived}} ───
    print(f"🔧 Replacing {{sampleReceived}} → '{received_date}'")
    replace_placeholders_in_docx(doc, [("{{sampleReceived}}", received_date)])

    # ─── 8) Replace {{sampleImage}} with image ───
    picture_ph = "{{sampleImage}}"
    pic_path = image_full_path if image_full_path.exists() else None
    if pic_path:
        print(f"🖼️ Will insert image from: {pic_path}")
    else:
        print(f"⚠️ No valid image to insert for {picture_ph}")

    # Paragraphs
    for p in doc.paragraphs:
        if replace_placeholder_across_runs(p.runs, picture_ph, pic_path, 50, 50):
            print(f"✅ Inserted {picture_ph} in a paragraph")

    # Tables
    for tbl in doc.tables:
        process_table(tbl, picture_ph, pic_path, 50, 50)

    print(f"🎉 Completed inserting submission data for sample {sample_id}")